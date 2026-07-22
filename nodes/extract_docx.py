from io import BytesIO
from PIL import Image
from docx import Document

from utils.ocr_utils import ocr_image


def _extract_table(table):
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _extract_docx_images(doc):
    image_text = ""
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            blob = rel.target_part.blob
            try:
                image_text += "\n[DOCX image OCR]\n"
                image_text += ocr_image(blob) + "\n"
            except Exception:
                continue
    return image_text


def extract_docx(state):
    doc = Document(state["file_path"])
    text = "\n=== DOCX Document ===\n"

    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"

    if doc.tables:
        text += "\n=== Tables ===\n"
        for table_index, table in enumerate(doc.tables, start=1):
            text += f"\n--- Table {table_index} ---\n"
            text += _extract_table(table) + "\n"

    image_text = _extract_docx_images(doc)
    if image_text.strip():
        text += "\n=== Images ===\n"
        text += image_text

    state["raw_text"] = text
    return state